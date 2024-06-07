from docx import Document as DocxDocument
from googletrans import Translator
from docx.shared import Pt
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

class FileTranslator:
    def __init__(self):
        self.translator = Translator()

    def extract_text_and_formatting_from_docx(self, file):
        doc = DocxDocument(file)
        content = []
        for i, paragraph in enumerate(doc.paragraphs):
            runs = []
            for j, run in enumerate(paragraph.runs):
                runs.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'font_name': run.font.name,
                })
            alignment = paragraph.alignment
            print(f"Paragraph {i}: Text: '{paragraph.text}', Alignment: {alignment}")
            content.append({'runs': runs, 'alignment': alignment})
        return content

    def translate_text(self, text, target_language):
        try:
            translated = self.translator.translate(text, dest=target_language)
            return translated.text
        except Exception as e:
            return f"Translation failed: {e}"

    def translate_runs(self, runs, target_language):
        translated_runs = []
        for run in runs:
            translated_text = self.translate_text(run['text'], target_language)
            translated_run = {
                'text': translated_text,
                'bold': run['bold'],
                'italic': run['italic'],
                'underline': run['underline'],
                'font_size': run['font_size'],
                'font_name': run['font_name'],
            }
            translated_runs.append(translated_run)
        return translated_runs

    def translate_docx_content(self, content, target_language):
        translated_content = []
        for paragraph in content:
            translated_paragraph = {
                'runs': self.translate_runs(paragraph['runs'], target_language),
                'alignment': paragraph['alignment']
            }
            translated_content.append(translated_paragraph)
        return translated_content

    def apply_formatting(self, p, run):
        r = p.add_run(run['text'])
        r.bold = run['bold']
        r.italic = run['italic']
        r.underline = run['underline']
        if run['font_size']:
            r.font.size = Pt(run['font_size'])
        if run['font_name']:
            r.font.name = run['font_name']

    def save_translated_docx(self, translated_content, file_name):
        doc = DocxDocument()
        
        for i, paragraph in enumerate(translated_content):
            p = doc.add_paragraph()
            for run in paragraph['runs']:
                self.apply_formatting(p, run)
            
            if paragraph['alignment'] is not None:
                print(f"Setting alignment for paragraph {i}: {paragraph['alignment']}")
                p.alignment = paragraph['alignment']
            else:
                print(f"No alignment for paragraph {i}")

        translated_file_path = f"{os.path.splitext(file_name)[0]}_translated.docx"
        doc.save(translated_file_path)
        return translated_file_path, "Translation and saving successful."